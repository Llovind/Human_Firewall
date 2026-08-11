import sys
import subprocess

try:
    import docx
except ImportError:
    print("python-docx belum terinstal. Mengunduh dan menginstal library...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # TITLE
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("IDE PENGEMBANGAN FITUR SOAR (AUTOMATED RESPONSE)\nMelengkapi Siklus Incident Response")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    # SUBTITLE
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Untuk: Daffa (Integration Engineer)\nDari: Tim SOC Developer")
    run_sub.font.size = Pt(10)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(99, 102, 241)

    # 1. KONSEP DASAR: MELENGKAPI SIKLUS INCIDENT RESPONSE
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    r1 = h1.add_run("1. KONSEP DASAR: MELENGKAPI SIKLUS INCIDENT RESPONSE")
    r1.font.size = Pt(13)
    r1.font.bold = True
    
    doc.add_paragraph(
        "Saat ini, siklus Incident Response di Human Firewall Lite masih berhenti di tahap 'Monitoring & Ticketing'. "
        "Artinya, sistem baru bisa mendeteksi ketika ada karyawan yang mengklik link berbahaya atau melaporkan email, "
        "kemudian mengirimkan peringatan ke analis SOC."
    )
    doc.add_paragraph(
        "Ide pengembangan selanjutnya adalah melengkapi siklus IR ini hingga tahap akhir, yaitu 'Automated Containment & Eradication' (Mitigasi Otomatis). "
        "Kita ingin mengubah platform ini menjadi sistem SOAR (Security Orchestration, Automation, and Response) sesungguhnya."
    )
    doc.add_paragraph(
        "Alur barunya akan seperti ini:\n"
        "1. Karyawan melaporkan email phishing dari Dashboard Webmail.\n"
        "2. Backend Flask akan mengekstrak Indicators of Compromise (IoC) seperti URL, IP Address, dan Email Pengirim.\n"
        "3. Flask mengirimkan data ancaman tersebut ke n8n melalui Webhook.\n"
        "4. n8n akan mengeksekusi otomatisasi mitigasi (misalnya memblokir IP atau menghapus email berbahaya) "
        "tanpa perlu intervensi manual dari analis."
    )

    # 2. ALTERNATIF INFRASTRUKTUR GRATIS / OPEN SOURCE
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    r2 = h2.add_run("2. IDE IMPLEMENTASI MITIGASI (GRATIS & OPEN SOURCE)")
    r2.font.size = Pt(13)
    r2.font.bold = True

    doc.add_paragraph(
        "Mengingat ini adalah tahap Proof of Concept, kita tidak akan menggunakan perangkat kelas enterprise yang mahal seperti Palo Alto atau Fortinet. "
        "Berikut adalah beberapa ide perangkat/layanan gratis dan open-source yang bisa kamu integrasikan di n8n untuk mengeksekusi payload mitigasi dari Flask:"
    )

    doc.add_paragraph(
        "A. Cloudflare WAF (Sangat Direkomendasikan)\n"
        "Cloudflare menyediakan versi gratis yang memiliki API lengkap. Melalui n8n, kita bisa menembak API Cloudflare "
        "untuk otomatis menambahkan IP hacker atau URL phishing ke dalam daftar hitam (blocklist) WAF."
    )

    doc.add_paragraph(
        "B. Pi-Hole / AdGuard Home (Level Jaringan Lokal)\n"
        "Bila pengujian dilakukan di lingkungan lokal, kita bisa menggunakan Pi-Hole (berbasis DNS sinkhole open-source). "
        "n8n dapat mengirimkan perintah API ke Pi-Hole untuk memasukkan domain jahat ke dalam blacklist DNS seketika."
    )

    doc.add_paragraph(
        "C. Eksekusi SSH ke Server Linux (iptables)\n"
        "Cara yang sangat brutal namun nyata: Menggunakan node SSH di n8n untuk masuk ke server Linux secara otomatis "
        "dan menjalankan perintah iptables (contoh: iptables -A INPUT -s <IP_HACKER> -j DROP) untuk memblokir IP."
    )

    doc.add_paragraph(
        "D. Microsoft 365 Developer Sandbox (Untuk Hapus Email)\n"
        "Untuk membuktikan kemampuan Search & Purge, kita dapat membuat akun Microsoft 365 Developer gratis. "
        "n8n dapat memanggil Microsoft Graph API untuk otomatis menghapus/karantina email yang dikirim dari alamat attacker di semua inbox pengguna."
    )

    # 3. KESIMPULAN
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(18)
    r3 = h3.add_run("3. TINDAK LANJUT")
    r3.font.size = Pt(13)
    r3.font.bold = True

    doc.add_paragraph(
        "Mohon pelajari alternatif di atas dan pilih salah satu (misalnya Cloudflare WAF atau Pi-Hole) yang menurut kamu "
        "paling memungkinkan untuk dieksekusi di n8n saat ini. Setelah diputuskan, tim Flask akan menyiapkan pengiriman webhook "
        "dengan struktur payload JSON yang disesuaikan."
    )

    output_path = r"C:\Human_Firewall\Ide_SOAR_Daffa.docx"
    doc.save(output_path)
    print(f"Laporan berhasil dibuat di: {output_path}")

if __name__ == "__main__":
    create_report()
