import os
import sys

try:
    import docx
except ImportError:
    print("python-docx belum terinstal. Mengunduh dan menginstal library...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # TITLE
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("SPESIFIKASI INTEGRASI SOAR (n8n)\nHuman Firewall Lite")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    # SUBTITLE
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Untuk: Daffa (Integration Engineer)\nDari: Tim SOC Developer")
    run_sub.font.size = Pt(10)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(99, 102, 241)

    # 1. KONSEP DASAR
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    r1 = h1.add_run("1. KONSEP DASAR (AUTOMATED RESPONSE)")
    r1.font.size = Pt(13)
    r1.font.bold = True
    
    doc.add_paragraph(
        "Platform Human Firewall Lite akan di-upgrade menjadi sistem SOAR (Security Orchestration, Automation, and Response). "
        "Saat ini, aplikasi Flask bertindak sebagai 'Otak' yang menerima laporan phishing dari karyawan. "
        "Tugas kamu di n8n adalah bertindak sebagai 'Otot' yang menerima webhook dari Flask dan melakukan mitigasi nyata di jaringan."
    )

    # 2. STRUKTUR PAYLOAD
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    r2 = h2.add_run("2. STRUKTUR PAYLOAD JSON DARI FLASK")
    r2.font.size = Pt(13)
    r2.font.bold = True

    doc.add_paragraph(
        "Ketika karyawan menekan tombol 'Lapor Phishing', Flask akan mem-parsing Indicators of Compromise (IoC) "
        "dari email tersebut dan mengirimkan HTTP POST Request ke URL Webhook n8n kamu. Payloadnya akan berbentuk seperti ini:"
    )

    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.5)
    run_code = p_code.add_run(
        "{\n"
        '  "event": "ioc_detected",\n'
        '  "reporter": "budi.santoso@perusahaan.com",\n'
        '  "confidence_score": "high",\n'
        '  "iocs": {\n'
        '    "malicious_urls": ["http://login-palsu-update.com"],\n'
        '    "sender_email": "admin@bca-update-palsu.com",\n'
        '    "source_ip": "103.45.67.89"\n'
        '  }\n'
        "}"
    )
    run_code.font.name = 'Courier New'
    run_code.font.size = Pt(9.5)

    # 3. TUGAS WORKFLOW
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(18)
    r3 = h3.add_run("3. TUGAS WORKFLOW DI n8n")
    r3.font.size = Pt(13)
    r3.font.bold = True

    doc.add_paragraph(
        "Kamu perlu membuat workflow baru di n8n dengan alur kerja berikut:\n\n"
        "1. Webhook Trigger:\n"
        "   Buat node Webhook (POST) untuk menerima payload JSON di atas. (Tolong berikan URL Webhook-nya ke tim Flask jika sudah jadi).\n\n"
        "2. Automated Purge (Microsoft 365 / Google Workspace):\n"
        "   Gunakan node HTTP Request atau node spesifik O365 untuk mencari email dengan 'sender_email' yang sama di semua inbox karyawan, lalu hapus secara paksa (Search & Purge).\n\n"
        "3. Network Blocking (Firewall):\n"
        "   Hubungkan n8n ke API Firewall perusahaan (Palo Alto / Fortinet) atau Cloudflare. Masukkan 'source_ip' dan 'malicious_urls' ke dalam Blocklist Policy secara otomatis.\n\n"
        "4. Callback ke SOC (Opsional):\n"
        "   Kirim balik HTTP POST ke API SOC kami (/api/soar-callback) untuk memberi tahu bahwa IP berhasil diblokir, sehingga Dashboard SOC bisa menampilkan log sukses."
    )

    output_path = r"C:\Human_Firewall\Task_SOAR_Daffa.docx"
    doc.save(output_path)
    print(f"Laporan berhasil dibuat di: {output_path}")

if __name__ == "__main__":
    create_report()
